import streamlit as st
import sys
import os
from pathlib import Path
import io
from datetime import datetime
from docx import Document

# Add the parent directory to path so we can import 'app'
sys.path.insert(0, str(Path(__file__).parent.parent))

from app.services.orchestration import OrchestrationService

# 1. Page Configuration
st.set_page_config(
    page_title="BankGuard - AI Compliance",
    page_icon="🏦",
    layout="wide"
)

# 2. Premium Design (CSS)
st.markdown("""
<style>
:root {
    --bg: #f4f7fb;
    --panel: #ffffff;
    --ink: #0f172a;
    --muted: #64748b;
    --brand: #0b1f3b;
    --accent: #b9892d;
}

html, body, [class*="stApp"] {
    background: linear-gradient(180deg, #f4f7fb 0%, #eef2f7 100%);
    color: var(--ink);
    font-family: "Segoe UI", "Montserrat", sans-serif;
}

.main { background: transparent; }

section[data-testid="stSidebar"] {
    background: #0b1f3b;
    border-right: 1px solid rgba(255, 255, 255, 0.08);
}

.stHeader {
    background: linear-gradient(135deg, #0b1f3b 0%, #153158 70%, #1f497d 100%);
    padding: 2rem 2rem;
    border-radius: 12px;
    color: #ffffff;
    margin-bottom: 16px;
    box-shadow: 0 10px 24px rgba(0,0,0,0.35);
}
.stHeader h1 { color: #ffffff !important; }
.stHeader p { color: #e5e7eb !important; font-size: 1.0rem; }

.badge {
    display: inline-block;
    font-size: 0.72rem;
    letter-spacing: 0.08em;
    text-transform: uppercase;
    padding: 0.3rem 0.6rem;
    border-radius: 999px;
    background: rgba(255, 255, 255, 0.15);
    border: 1px solid rgba(255, 255, 255, 0.3);
    margin-bottom: 0.6rem;
}

.card {
    background: #ffffff;
    color: #0f172a;
    padding: 1.1rem;
    border-radius: 10px;
    box-shadow: 0 8px 18px rgba(0,0,0,0.2);
    border: 1px solid #e5e7eb;
    margin-bottom: 0.9rem;
}

.metric-box {
    background: #f8fafc;
    color: #0f172a;
    padding: 0.9rem;
    border-radius: 10px;
    text-align: center;
    border: 1px solid #e2e8f0;
    box-shadow: 0 6px 14px rgba(0,0,0,0.18);
}

.priority-high {
    color: #991b1b;
    font-weight: 700;
    background: #fef2f2;
    padding: 2px 8px;
    border-radius: 4px;
}
.priority-medium {
    color: #92400e;
    font-weight: 700;
    background: #fffbeb;
    padding: 2px 8px;
    border-radius: 4px;
}

.stTabs [data-baseweb="tab"] {
    color: #1f2937 !important;
    font-weight: 600;
}
.stTabs [data-baseweb="tab"][aria-selected="true"] {
    color: #0b1f3b !important;
}

.stButton > button {
    background: linear-gradient(135deg, #0b1f3b 0%, #153158 100%);
    color: #ffffff !important;
    border: none;
    padding: 0.55rem 1.1rem;
    border-radius: 8px;
    font-weight: 600;
}
</style>
""", unsafe_allow_html=True)







def build_docx_report(res):
    doc = Document()
    doc.add_heading("BankGuard Compliance Report", level=1)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    summary = res.get("summary", {})
    validation = res.get("validation", {})
    impact = res.get("impact", {})

    doc.add_heading("Executive Summary", level=2)
    doc.add_paragraph(summary.get("summary_text", "No summary available."))

    doc.add_heading("Key Metrics", level=2)
    doc.add_paragraph(f"Obligations Found: {len(summary.get('obligations', []))}")
    doc.add_paragraph(f"Deadlines Extracted: {len(summary.get('deadlines', []))}")
    doc.add_paragraph(f"Validation Status: {validation.get('validation_status', 'N/A')}")

    doc.add_heading("Key Obligations", level=2)
    obligations = summary.get("obligations", [])
    if obligations:
        for ob in obligations:
            doc.add_paragraph(ob, style="List Bullet")
    else:
        doc.add_paragraph("No obligations detected.")

    doc.add_heading("Deadlines", level=2)
    deadlines = summary.get("deadlines", [])
    if deadlines:
        for dl in deadlines:
            doc.add_paragraph(dl, style="List Bullet")
    else:
        doc.add_paragraph("No deadlines detected.")

    doc.add_heading("Impact Mapping", level=2)
    if impact.get("departments"):
        doc.add_paragraph("Departments: " + ", ".join(impact.get("departments", [])))
    else:
        doc.add_paragraph("Departments: N/A")
    if impact.get("products"):
        doc.add_paragraph("Products: " + ", ".join(impact.get("products", [])))
    else:
        doc.add_paragraph("Products: N/A")

    doc.add_heading("Action Plan", level=2)
    action_plan = res.get("action_plan", [])
    if action_plan:
        for action in action_plan:
            owner = action.get("responsible_department", "")
            priority = action.get("priority", "")
            line = f"{action.get('action','')} (Owner: {owner}, Priority: {priority})"
            doc.add_paragraph(line, style="List Bullet")
    else:
        doc.add_paragraph("No actions generated.")

    doc.add_heading("Evidence Citations", level=2)
    evidence = validation.get("evidence", [])
    if evidence:
        for ev in evidence:
            doc.add_paragraph(f"Obligation: {ev.get('obligation','')}", style="List Bullet")
            snippets = ev.get("evidence_snippets", [])
            if snippets:
                snip = snippets[0]
                doc.add_paragraph(f"Page: {snip.get('page','N/A')}")
                doc.add_paragraph(snip.get("text", ""))
    else:
        doc.add_paragraph("No evidence available.")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# 3. Initialization
if 'orchestrator' not in st.session_state:
    st.session_state.orchestrator = OrchestrationService()
if 'result' not in st.session_state:
    st.session_state.result = None

# Header
st.markdown(
    '<div class="stHeader">'
    '<div class="badge">RegTech • RBI/FEMA</div>'
    '<h1 style="color:#ffffff;">🏦 BankGuard</h1>'
    '<p style="color:#ffffff;">Agentic AI Compliance Platform for RBI & FEMA</p>'
    '<p style="color:#eaf2ff; font-size:0.95rem; margin-top:0.35rem;">'
    'Cut analysis time, reduce missed obligations, and align internal policy with every regulatory update.'
    '</p>'
    '</div>',
    unsafe_allow_html=True,
)

# 4. Sidebar Controls
with st.sidebar:
    st.subheader("Private Policy Repository")
    policy_files = st.file_uploader(
        "Upload internal policy PDFs",
        type=["pdf"],
        accept_multiple_files=True,
    )
    if st.button("Add Policies"):
        if policy_files:
            with st.spinner("Adding internal policies to private repository..."):
                added = 0
                for policy_file in policy_files:
                    st.session_state.orchestrator.add_internal_policy(
                        policy_file.read(),
                        policy_file.name,
                    )
                    added += 1
            st.success(f"Added {added} internal policy file(s).")
        else:
            st.warning("Please upload at least one policy PDF.")
    st.header("⚙️ Controls")
    with st.expander("Database Management", expanded=False):
        st.write("Clear all stored documents and embeddings.")
        if st.button("Clear Database"):
            st.session_state.orchestrator.clear_database()
            st.success("Database Cleared")

# 5. Main Tabs
tab1, tab2, tab3 = st.tabs(["📄 Document Analysis", "📊 Results Dashboard", "ℹ️ About"])

with tab1:
    st.subheader("Upload Regulatory Document")
    uploaded_file = st.file_uploader("Select RBI/FEMA PDF", type=['pdf'])
    
    if st.button("🚀 Analyze Document", type="primary"):
        if uploaded_file:
            with st.spinner("Pipeline running: Ingesting → RAG → 4 Agents..."):
                file_bytes = uploaded_file.read()
                res = st.session_state.orchestrator.process_pdf(file_bytes, uploaded_file.name)
                st.session_state.result = res
                st.success("Analysis Complete!")
                st.toast("Analysis completed successfully.", icon="✅")
        else:
            st.warning("Please upload a file first.")

with tab2:
    if st.session_state.result:
        res = st.session_state.result

        # Summary Row
        col1, col2, col3 = st.columns(3)
        with col1:
            st.markdown(
                f'<div class="metric-box"><h3>{len(res["summary"]["obligations"])}</h3><p>Obligations Found</p></div>',
                unsafe_allow_html=True,
            )
        with col2:
            st.markdown(
                f'<div class="metric-box"><h3>{len(res["summary"]["deadlines"])}</h3><p>Deadlines Extracted</p></div>',
                unsafe_allow_html=True,
            )
        with col3:
            st.markdown(
                f'<div class="metric-box"><h3>{res["validation"]["validation_status"]}</h3><p>Validation Status</p></div>',
                unsafe_allow_html=True,
            )

        st.markdown("---")

        st.subheader("Executive Summary")
        report_buf = build_docx_report(res)
        st.download_button(
            label="Download Report (DOCX)",
            data=report_buf.getvalue(),
            file_name=f"bankguard_report_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        st.markdown(
            f'<div class="card">{res["summary"].get("summary_text", "No summary available.")}</div>',
            unsafe_allow_html=True,
        )

        c_left, c_right = st.columns(2)
        with c_left:
            st.subheader("Key Obligations")
            obligations = res["summary"].get("obligations", [])
            if obligations:
                for i, ob in enumerate(obligations[:8], start=1):
                    st.markdown(f"{i}. {ob}")
            else:
                st.info("No obligations detected.")

            st.subheader("Deadlines")
            deadlines = res["summary"].get("deadlines", [])
            if deadlines:
                for i, dl in enumerate(deadlines[:8], start=1):
                    st.markdown(f"{i}. {dl}")
            else:
                st.info("No deadlines detected.")

        with c_right:
            st.subheader("Impact Mapping")
            depts = res["impact"].get("departments", [])
            products = res["impact"].get("products", [])
            st.write(f"**Departments:** {', '.join(depts) if depts else 'N/A'}")
            st.write(f"**Products:** {', '.join(products) if products else 'N/A'}")

            st.subheader("Action Plan")
            if res.get("action_plan"):
                st.table(
                    [
                        {
                            "Action": a.get("action", ""),
                            "Owner": a.get("responsible_department", ""),
                            "Priority": a.get("priority", ""),
                        }
                        for a in res["action_plan"]
                    ]
                )
            else:
                st.info("No actions generated.")

        st.subheader("Evidence Citations")
        if res["validation"].get("evidence"):
            for ev in res["validation"]["evidence"]:
                page = "N/A"
                if ev.get("evidence_snippets"):
                    page = ev["evidence_snippets"][0].get("page", "N/A")
                with st.expander(f"Obligation: {ev['obligation'][:80]}"):
                    st.caption(f"Page: {page}")
                    if ev.get("evidence_snippets"):
                        st.write(ev["evidence_snippets"][0].get("text", ""))
        else:
            st.info("No evidence available.")

    else:
        st.info("No analysis found. Upload a document in the first tab.")

with tab3:
    st.markdown("""
    ### How BankGuard Works
    1. **PDF Ingestion**: Parses regulatory text with page markers.
    2. **RAG Retrieval**: Searches for relevant chunks using Chroma Vector DB & HuggingFace.
    3. **Agentic Pipeline**: 
       * **Summarizer**: Extracts hard obligations and dates.
       * **Impact Mapper**: Links keywords to banking departments.
       * **Validator**: Cross-references extraction against source text for citations.
       * **Action Planner**: Generates prioritized tasks.
    """)
